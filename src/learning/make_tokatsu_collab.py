# -*- coding: utf-8 -*-
"""트릭컬 × ド葛本社(니지산지) 콜라보 실장 가능 항목 분석 — 한중 양어 Word 생성
형식: 질문 4번에 대한 답변 문서 (a)~ 항목별
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build(lang):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)

    def kv(k, v):
        p = doc.add_paragraph()
        r = p.add_run(f'[{k}] ')
        r.bold = True
        p.add_run(v)

    if lang == 'kr':
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run('트릭컬 리바이브 × ド葛本社 (니지산지)\n콜라보 시점별 게임 내 실장 가능 항목')
        r.bold = True; r.font.size = Pt(16)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('질문 4. 虹彩虹社(니지산지) ド葛本社 VTuber 그룹과 콜라보 시,\n각 콜라보 시점(노드)에 대응하는 게임 내 실장 가능 콜라보 콘텐츠').font.color.rgb = RGBColor(0x66,0x66,0x66)

        doc.add_paragraph()
        d = doc.add_paragraph()
        d.add_run('※ 기본 전제 — 아래 내용은 트릭컬 리바이브의 기존 게임 내 시스템(사도/카드/사복/가구/펫/명찰/프레임/로비/보이스)을 기준으로, ド葛本社 콜라보 시 각 시점에서 실장 가능한 콘텐츠를 정리한 것이다. 구체적 계약 조건·일러스트 확보 여부에 따라 변동될 수 있음.').font.size = Pt(9)

        doc.add_heading('a) 콜라보 전 사전 단계 (티저 기간)', level=2)
        kv('실장 항목', '로비(大厅) 배경 교체, 로딩 화면, 예고 팝업, 사전등록/출석 이벤트')
        kv('상세', '콜라보 티저 이미지를 로비 배경 및 로딩에 적용. "謎の来电" 형식의 티저 이벤트(전화기를 모티브로 한 트릭컬 특성과 ド葛本社 멤버의 호러/괴기 콘텐츠 친화성 연결) 가능. SNS 예고 → 인게임 팝업 유도 구조.')
        kv('참고', '현재 holo 측과도 콜라보를 검토 중으로, 구체적 형태는 확정 전 — 본 안은 ド葛本社 기준.')

        doc.add_heading('b) 콜라보 시작 (1차 업데이트)', level=2)
        kv('캐릭터', 'ド葛本社 멤버 4인(葛葉/ドーラ/本間ひまわり/社築)을 사도(使者)로 실장 — 카드 획득 이벤트/뽑기')
        kv('스킨', '멤버 상징 컨셉 사복(私服) 4종 — 각 캐릭터 시그니처 컬러/모티프 반영')
        kv('보이스', '멤버 실제 보이스를 활용한 캐릭터 보이스(인사/터치/볼당기기 대응) ※음원 라이선스 계약 필요')
        kv('이벤트', '콜라보 전용 테마극장(스토리+미니게임), 콜라보 로그인 보상')

        doc.add_heading('c) 콜라보 중반 (2차 업데이트)', level=2)
        kv('펫', '멤버 상징 펫 4종 (예: 葛葉=까마귀/박쥐 모티프, ひまわり=해바라기 모티프 등) — 테마극장/미니카 시장 연동 획득')
        kv('가구', '콜라보 테마 가구 세트 (멤버 방 컨셉 재현 — 게임 내 가구 시스템 활용)')
        kv('명찰', '멤버 이름표/명찰 장식 4종')
        kv('프레임', '콜라보 기념 프로필 프레임 (개인/단체 버전)')

        doc.add_heading('d) 콜라보 후반 (라이브 연동)', level=2)
        kv('실시간 연동', '1주년 생방송 등 라이브 방송과 인게임 연동 — 시청 특전 코드, 방송 시청자 투표 → 인게임 반영')
        kv('아이콘', '콜라보 한정 아이콘(VG콜라보 4종 선례와 동일 구조) — 이벤트 획득')
        kv('로비 배경', '콜라보 기념 로비 배경 (영구/기간한정)')

        doc.add_heading('e) 콜라보 종료 (회수 처리)', level=2)
        kv('전환 안내', '콜라보 콘텐츠 종료 사전 안내(공지) — VG콜라보 아이콘 사례처럼 종료 후 획득 방식 별도 안내 필수')
        kv('보존', '획득한 스킨/가구/펫/프레임은 영구 보유, 이벤트 재화는 골드 자동 전환(기존 규칙)')

        doc.add_heading('요약표', level=2)
        table = doc.add_table(rows=6, cols=5)
        table.style = 'Light Grid Accent 1'
        heads = ['콜라보 시점', '캐릭터/스킨', '가구/펫', '명찰/프레임', '로비/기타']
        for i, h in enumerate(heads):
            c = table.rows[0].cells[i]; c.text = h
            for pp in c.paragraphs:
                for rr in pp.runs: rr.bold = True
        rows = [
            ('a) 사전 티저', '-', '-', '-', '로비 배경, 팝업, 사전이벤트'),
            ('b) 1차 업데이트', '사도 4인 + 사복 4종 + 보이스', '-', '-', '테마극장, 로그인 보상'),
            ('c) 2차 업데이트', '-', '가구 세트 + 펫 4종', '명찰 4종 + 프레임', '-'),
            ('d) 라이브 연동', '-', '-', '아이콘(한정)', '실시간 코드, 로비 배경(기념)'),
            ('e) 종료', '-', '-', '-', '종료 안내, 재료 전환'),
        ]
        for i, row in enumerate(rows, start=1):
            for j, v in enumerate(row):
                table.rows[i].cells[j].text = v

        p = doc.add_paragraph()
        p.add_run('\n※ holo 콜라보 관련: 현재 검토 중으로 구체적 형태 확정 전 (질문 양식 참조).').font.size = Pt(9)
        out = r'C:\Users\user\Desktop\트릭컬_도카츠혼샤_콜라보_실장가능항목_한국어.docx'

    else:  # 중국어
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run('트릭컬 리바이브 × ド葛本社(彩虹社/니지산지)\n联动各阶段游戏内可实装项目')
        r.bold = True; r.font.size = Pt(16)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('问题4. 如果与彩虹社【ド葛本社】这个VTuber组合合作,\n以下合作节点对应的游戏内可以支持实装的联动内容').font.color.rgb = RGBColor(0x66,0x66,0x66)

        doc.add_paragraph()
        d = doc.add_paragraph()
        d.add_run('※ 基本前提 — 以下内容以트릭컬 리바이브(Trickcal Revive)现有游戏内系统(使徒/卡牌/私服/家具/宠物/名牌/头像框/大厅/语音)为基础, 整理各合作节点可实装的联动内容。具体合同条件及插图授权情况可能导致变动。').font.size = Pt(9)

        doc.add_heading('a) 合作前预热期(预热期间)', level=2)
        kv('实装项目', '大厅背景更换、加载画面、预告弹窗、预约/签到活动')
        kv('详情', '将联动预热图应用于大厅背景及加载画面。可利用트릭컬特性和ド葛本社成员的恐怖/怪奇内容亲和性, 打造"神秘来电"形式的预热活动。SNS预告→游戏内弹窗引导结构。')
        kv('备注', '目前也在与holo合作讨论中, 具体合作形式确认中 — 本方案以ド葛本社为准。')

        doc.add_heading('b) 合作开始(第一次更新)', level=2)
        kv('角色', '将ド葛本社成员4人(葛叶/ドーラ/本间ひまわり/社筑)以使徒形式实装 — 卡牌获取活动/抽卡')
        kv('皮肤', '成员象征概念私服4种 — 反映各角色标志色/主题元素')
        kv('语音', '利用成员真实语音的角色语音(问候/触摸/捏脸对应) ※需语音版权合同')
        kv('活动', '联动专用主题剧场(剧情+小游戏)、联动登录奖励')

        doc.add_heading('c) 合作中期(第二次更新)', level=2)
        kv('宠物', '成员象征宠物4种 (例: 葛叶=乌鸦/蝙蝠主题, ひまわり=向日葵主题等) — 主题剧场/迷你车市场联动获取')
        kv('家具', '联动主题家具套装(再现成员房间概念 — 运用游戏内家具系统)')
        kv('名牌', '成员名牌/名牌装饰4种')
        kv('头像框', '联动纪念头像框(个人/团体版本)')

        doc.add_heading('d) 合作后期(直播联动)', level=2)
        kv('实时联动', '周年直播等与游戏内联动 — 观看特典兑换码、直播观众投票→游戏内反映')
        kv('图标', '联动限定图标(VG联动4种先例相同结构) — 活动获取')
        kv('大厅背景', '联动纪念大厅背景(永久/限时)')

        doc.add_heading('e) 合作结束(回收处理)', level=2)
        kv('转换告知', '联动内容结束提前告知(公告) — 如VG联动图标事例, 结束后获取方式需另行告知')
        kv('保留', '已获得的皮肤/家具/宠物/头像框永久保留, 活动货币自动转换为金币(现有规则)')

        doc.add_heading('汇总表', level=2)
        table = doc.add_table(rows=6, cols=5)
        table.style = 'Light Grid Accent 1'
        heads = ['合作节点', '角色/皮肤', '家具/宠物', '名牌/头像框', '大厅/其他']
        for i, h in enumerate(heads):
            c = table.rows[0].cells[i]; c.text = h
            for pp in c.paragraphs:
                for rr in pp.runs: rr.bold = True
        rows = [
            ('a) 预热期', '-', '-', '-', '大厅背景、弹窗、预热活动'),
            ('b) 第一次更新', '使徒4人+私服4种+语音', '-', '-', '主题剧场、登录奖励'),
            ('c) 第二次更新', '-', '家具套装+宠物4种', '名牌4种+头像框', '-'),
            ('d) 直播联动', '-', '-', '图标(限定)', '实时兑换码、大厅背景(纪念)'),
            ('e) 结束', '-', '-', '-', '结束告知、材料转换'),
        ]
        for i, row in enumerate(rows, start=1):
            for j, v in enumerate(row):
                table.rows[i].cells[j].text = v

        p = doc.add_paragraph()
        p.add_run('\n※ holo联动相关: 目前讨论中, 具体形式确认中(参照问题格式)。').font.size = Pt(9)
        out = r'C:\Users\user\Desktop\트릭컬_도카츠혼샤_콜라보_실장가능항목_中文版.docx'

    doc.save(out)
    return out

if __name__ == '__main__':
    print(build('kr'))
    print(build('cn'))
