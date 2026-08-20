# -*- coding: utf-8 -*-
"""트릭컬 1주년 팝업스토어 이스터에그 기획서 — 한국어 번역 Word 생성"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 스타일
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def para(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    return p

def kv(k, v):
    p = doc.add_paragraph()
    r1 = p.add_run(f'[{k}] ')
    r1.bold = True
    p.add_run(v)
    return p

# ── 제목 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('『트릭컬』 1주년 팝업스토어\n체험형 이스터에그 시책 기획서')
r.bold = True
r.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('—— 「이세계에서 온 전화기 (가칭)」')
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── 1. 기획 배경·목적 ──
h1('1. 기획 배경 및 목적')
kv('배경', '1주년 생방송의 "호러/괴기 테마" 프로모션과 연동하여, 팝업스토어 매장 내에 미스터리하면서도 인터랙티브한 체험형 장치를 설치합니다.')
kv('목적', '내방객(유저)의 몰입감과 탐험 체험을 높이고 캐릭터와의 친밀감을 깊게 하는 동시에, 온라인 1주년 생방송으로의 유도(티저 효과)를 도모합니다.')

# ── 2. 장치 콘셉트 ──
h1('2. 장치 콘셉트 및 외관 이미지')
kv('장치 외형', '레트로 버튼식 전화기 (또는 녹음기기).')
kv('외관·VMD 연출', '')
para('전화기 옆에 "경고문"이나 "수수께끼의 실험 메모/괴기 현상 기록 메모지"를 배치 (푸시버튼을 누르도록 유저에게 유도)')
para('1주년 테마에 맞춘 테이블/벽면 장식')
kv('체험 방법', '유저가 수화기를 들고 키패드(0~9)를 누르면, 수화기에서 각 버튼에 대응하는 캐릭터의 "수수께끼의 음성/호러풍 보이스"(노이즈 섞인 음성메시지, SOS 신호, 생방송 예고 암호 등)가 재생됩니다.')

# ── 3. 인터랙티브 & 음성 로직 ──
h1('3. 인터랙티브 및 음성 로직')

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, t in enumerate(['조작/트리거', '음성 내용안', '연출 효과 / 목적']):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ('수화기를 든다 (대기 상태)', '희미한 노이즈 / 화이트노이즈 / 으스스한 환경음', '불길한 몰입감 연출'),
    ('버튼 1~6', '메인/인기 캐릭터의 1주년 테마 보이스\n(단속적인 통신음, 괴기한 메시지 등)', '캐릭터와의 인터랙티브 체험, 팬 만족도 향상'),
    ('버튼 7~9', '미지의 캐릭터 보이스 / 으스스한 괴음', '호러 분위기를 높이고 SNS 화제화 겨냥'),
    ('버튼 0', '1주년 생방송 티저 보이스\n(예: "○월 ○일… 방송을 봐라… 무서운 일이 일어난다…" 등)', '온라인 1주년 생방송 유도·예고 효과'),
    ('', '', ''),
]
# 마지막 빈 행 제거
table._tbl.remove(table.rows[5]._tr)

for i, (a, b, c) in enumerate(rows_data[:4], start=1):
    cells = table.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

# ── 6. 협조 요청 ──
h1('6. 팝업스토어 측 확인 및 협조 요청 사항 (Checklist)')
kv('설치 영역', '매장 내 코너 또는 벽 쪽 (약 0.5㎡~1㎡ 정도)')
kv('기기 조달·설치', '본 개조 전화기는 당사에서 완제품을 납품·설치하는 방식이나, 귀사 시공팀이 기기 조달·고정 보조를 맡아주시는 방식 중 협의 후 결정하고자 합니다.')
kv('VMD/내부 인테리어', '전화기 주변 테이블·배경 벽면 장식에 대해, 팝업스토어 전체 인테리어 디자인·VMD 계획에 포함해 주시는 것이 가능할지요?')
kv('음성 데이터 납품 예정', '테스트용 조정 완료된 MP3 데이터는 [8월 31일]까지 납품할 예정입니다.')

doc.save(r'C:\Users\user\Desktop\트릭컬_1주년_팝업스토어_이스터에그_기획서_한국어.docx')
print('저장 완료')
